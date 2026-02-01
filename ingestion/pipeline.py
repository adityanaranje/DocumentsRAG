import os
import re
import pdfplumber
import docx
import warnings
import logging
from typing import List, Tuple
from pathlib import Path
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from ingestion.metadata import MetadataExtractor
from tqdm import tqdm

# Suppress PDF extraction warnings that clutter the terminal
logging.getLogger("pdfminer").setLevel(logging.ERROR)
warnings.filterwarnings("ignore", message=".*Cannot set gray stroke color.*")


class IngestionPipeline:
    """
    Enhanced ingestion pipeline with:
    - Differentiated chunk sizes for brochures vs CIS
    - Section-aware chunking with semantic header detection
    - Micro-chunking for tables
    """
    
    def __init__(self, docs_dir: str):
        self.docs_dir = docs_dir
        self.metadata_extractor = MetadataExtractor(docs_dir)
        
        # Brochure: 500-800 tokens (~650 avg), larger chunks for narrative content
        self.brochure_splitter = RecursiveCharacterTextSplitter(
            chunk_size=2600,  # ~650 tokens * 4 chars/token
            chunk_overlap=400,  # ~100 tokens overlap
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # CIS: 250-400 tokens (~325 avg), smaller for dense factual content
        self.cis_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1300,  # ~325 tokens * 4 chars/token
            chunk_overlap=160,  # ~40 tokens overlap
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        # Table micro-chunker: Very small for precise retrieval
        self.table_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,  # ~200 tokens for tables
            chunk_overlap=100,
            separators=["\n", "|", " ", ""]
        )

    def _extract_with_tables(self, file_path: str) -> List[Tuple[Document, str]]:
        """
        Uses pdfplumber to extract text while preserving table layouts.
        Returns list of (Document, chunk_type) tuples.
        """
        documents = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                
                # Add regular text content
                if text.strip():
                    documents.append((
                        Document(page_content=text, metadata={"page": i + 1}),
                        "text"
                    ))
                
                # Check for tables and add as separate micro-chunks
                tables = page.extract_tables()
                if tables:
                    for table_idx, table in enumerate(tables):
                        table_str = self._format_table(table)
                        if table_str.strip():
                            documents.append((
                                Document(
                                    page_content=table_str,
                                    metadata={"page": i + 1, "table_index": table_idx}
                                ),
                                "table"
                            ))
        return documents

    def _format_table(self, table: List[List]) -> str:
        """Format table rows into markdown-style pipe-separated format."""
        rows = []
        for row in table:
            clean_row = [str(cell).replace("\n", " ").strip() if cell else "" for cell in row]
            rows.append("| " + " | ".join(clean_row) + " |")
        return "\n".join(rows)

    def _extract_from_docx(self, file_path: str) -> List[Tuple[Document, str]]:
        """
        Extracts text from DOCX files with table separation.
        """
        doc = docx.Document(file_path)
        documents = []
        
        # Extract paragraphs as text
        para_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                para_text.append(para.text)
        
        if para_text:
            documents.append((
                Document(page_content="\n".join(para_text), metadata={"page": 1}),
                "text"
            ))
        
        # Extract tables separately
        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            
            table_str = "\n".join(rows)
            if table_str.strip():
                documents.append((
                    Document(
                        page_content=table_str,
                        metadata={"page": 1, "table_index": table_idx}
                    ),
                    "table"
                ))
        
        return documents

    def _get_splitter_for_doc_type(self, doc_type: str):
        """Return appropriate splitter based on document type."""
        if doc_type == "cis":
            return self.cis_splitter
        return self.brochure_splitter

    def _chunk_with_metadata(self, raw_docs: List[Tuple[Document, str]], 
                             file_metadata: dict) -> List[Document]:
        """
        Split documents with proper metadata including section and chunk_type.
        """
        all_chunks = []
        doc_type = file_metadata.get("document_type", "brochure")
        text_splitter = self._get_splitter_for_doc_type(doc_type)
        
        for doc, chunk_type in raw_docs:
            # Choose splitter based on content type
            if chunk_type == "table":
                splitter = self.table_splitter
            else:
                splitter = text_splitter
            
            # Split the document
            chunks = splitter.split_documents([doc])
            
            # Enrich each chunk with metadata
            for chunk in chunks:
                # Detect section from content
                section = MetadataExtractor.detect_section(chunk.page_content)
                
                # Merge file-level metadata
                chunk.metadata.update(file_metadata)
                chunk.metadata["chunk_type"] = chunk_type
                chunk.metadata["section"] = section
                
                all_chunks.append(chunk)
        
        return all_chunks

    def load_and_process_documents(self) -> List[Document]:
        """Process all documents in the docs directory."""
        all_chunks = []
        files_to_process = []

        for root, _, files in os.walk(self.docs_dir):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx')):
                    files_to_process.append(os.path.join(root, file))

        print(f"Found {len(files_to_process)} documents for ingestion.")

        for file_path in tqdm(files_to_process, desc="Processing"):
            try:
                file_metadata = self.metadata_extractor.extract_from_path(file_path)
                
                # Load based on file extension
                if file_path.lower().endswith('.pdf'):
                    raw_docs = self._extract_with_tables(file_path)
                elif file_path.lower().endswith('.docx'):
                    raw_docs = self._extract_from_docx(file_path)
                else:
                    continue
                
                chunks = self._chunk_with_metadata(raw_docs, file_metadata)
                all_chunks.extend(chunks)
                
            except Exception as e:
                print(f"Error processing {file_path}: {str(e)}")

        print(f"Total chunks created: {len(all_chunks)}")
        return all_chunks

    def process_single_file(self, file_path: str) -> List[Document]:
        """
        Process a single file for incremental ingestion.
        Used by the Flask upload endpoint.
        """
        file_metadata = self.metadata_extractor.extract_from_path(file_path)
        
        if file_path.lower().endswith('.pdf'):
            raw_docs = self._extract_with_tables(file_path)
        elif file_path.lower().endswith('.docx'):
            raw_docs = self._extract_from_docx(file_path)
        else:
            return []
        
        return self._chunk_with_metadata(raw_docs, file_metadata)


